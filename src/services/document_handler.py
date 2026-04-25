# src/services/document_handler.py

from typing import Dict, List
import re
from io import BytesIO


class DocumentHandler:
    """
    Handles document parsing for PDF, DOCX, and TXT files.
    """

    @staticmethod
    def extract_text_from_txt(content: bytes) -> Dict:
        """Extract text from plain text file."""
        try:
            text = content.decode('utf-8').strip()
            return {
                "success": True,
                "text": text,
                "error": None
            }
        except UnicodeDecodeError:
            try:
                text = content.decode('latin-1').strip()
                return {
                    "success": True,
                    "text": text,
                    "error": None
                }
            except:
                return {
                    "success": False,
                    "text": "",
                    "error": "Unable to decode text file"
                }

    @staticmethod
    def extract_text_from_pdf(content: bytes) -> Dict:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text

            if not text.strip():
                return {
                    "success": False,
                    "text": "",
                    "error": "PDF contains no extractable text (possibly scanned image)"
                }

            return {
                "success": True,
                "text": text.strip(),
                "error": None,
                "total_pages": len(pdf_reader.pages)
            }

        except ImportError:
            return {
                "success": False,
                "text": "",
                "error": "PyPDF2 library not installed. Install with: pip install PyPDF2"
            }
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "error": f"Error reading PDF: {str(e)}"
            }

    @staticmethod
    def extract_text_from_docx(content: bytes) -> Dict:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            docx_file = BytesIO(content)
            doc = Document(docx_file)
            
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            if not text.strip():
                return {
                    "success": False,
                    "text": "",
                    "error": "DOCX file contains no text"
                }

            return {
                "success": True,
                "text": text.strip(),
                "error": None,
                "total_paragraphs": len(doc.paragraphs)
            }

        except ImportError:
            return {
                "success": False,
                "text": "",
                "error": "python-docx library not installed. Install with: pip install python-docx"
            }
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "error": f"Error reading DOCX: {str(e)}"
            }

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text."""
        # Remove extra whitespace
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        return text

    @staticmethod
    def chunk_text(text: str, max_chunk_size: int = 1500, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks for analysis.
        
        Args:
            text: Text to chunk
            max_chunk_size: Maximum characters per chunk
            overlap: Number of overlapping characters between chunks
        """
        if len(text) <= max_chunk_size:
            return [text]

        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= max_chunk_size:
                current_chunk += sentence + " "
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())

                # Create overlap
                if len(chunks) > 0:
                    # Keep last few sentences for overlap
                    last_sentences = " ".join(sentence.split()[-20:])
                    current_chunk = last_sentences + " " + sentence + " "
                else:
                    current_chunk = sentence + " "

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks if chunks else [text]

    @staticmethod
    def create_summary(text: str, max_length: int = 500) -> str:
        """Create a summary of the text."""
        if len(text) <= max_length:
            return text

        # Take first max_length characters, ending at sentence boundary
        summary = text[:max_length]
        last_period = summary.rfind('.')
        
        if last_period > max_length * 0.7:  # If period is reasonably close
            summary = summary[:last_period + 1]
        else:
            summary = summary.rsplit(' ', 1)[0] + "..."

        return summary


# ============== MAIN FUNCTION ============== #

def extract_text_from_document(file_content: bytes, file_extension: str) -> Dict:
    """
    Extract text from any supported document format.

    Args:
        file_content: Raw file bytes
        file_extension: File extension (pdf, docx, txt)

    Returns:
        Dict with:
            - success: bool
            - text: str (full text)
            - summary: str (first 500 chars)
            - chunks: List[str] (for long documents)
            - error: str (if failed)
            - metadata: dict
    """

    file_extension = file_extension.lower()

    # Route to appropriate handler
    if file_extension == 'pdf':
        result = DocumentHandler.extract_text_from_pdf(file_content)
    elif file_extension == 'docx' or file_extension == 'doc':
        result = DocumentHandler.extract_text_from_docx(file_content)
    elif file_extension == 'txt':
        result = DocumentHandler.extract_text_from_txt(file_content)
    else:
        return {
            "success": False,
            "error": f"Unsupported file type: {file_extension}",
            "text": "",
            "chunks": []
        }

    if not result["success"]:
        return {
            "success": False,
            "error": result.get("error", "Unknown error"),
            "text": "",
            "chunks": []
        }

    text = DocumentHandler.clean_text(result["text"])

    # Create chunks if text is long
    if len(text) > 2000:
        chunks = DocumentHandler.chunk_text(text, max_chunk_size=1500)
    else:
        chunks = [text]

    summary = DocumentHandler.create_summary(text)

    return {
        "success": True,
        "text": text,
        "summary": summary,
        "chunks": chunks,
        "metadata": {
            "file_extension": file_extension,
            "total_length": len(text),
            "total_chunks": len(chunks),
            "total_pages": result.get("total_pages", 1),
            "total_paragraphs": result.get("total_paragraphs", 1)
        }
    }


# ============== TEST ============== #

if __name__ == "__main__":
    # Example usage
    print("Document Handler Module Loaded Successfully")
    print("Supported formats: PDF, DOCX, TXT")
    print("Use extract_text_from_document(file_content, file_extension) to extract text")